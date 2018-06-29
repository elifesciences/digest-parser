from docx import Document


def parse_content(file_name):
    "return all the content for testing"
    content = ''
    document = Document(file_name)
    # print(document.core_properties.author)
    # print(document.paragraphs)
    for para in document.paragraphs:
        content += join_runs(para.runs) + "\n"
    return content


def html_open_tag(style):
    "for the style return the HTML open tag"
    style_map = {
        'italic': '<i>',
        'bold': '<b>',
        'subscript': '<sub>',
        'superscript': '<sup>'
    }
    return style_map.get(style)


def html_close_tag(style):
    "for the style return the HTML close tag"
    style_map = {
        'italic': '</i>',
        'bold': '</b>',
        'subscript': '</sub>',
        'superscript': '</sup>'
    }
    return style_map.get(style)


def html_open_close_tag(style):
    "return the HTML open and close tags for the style"
    return html_open_tag(style), html_close_tag(style)


def run_contains_break(run):
    "check if a document run contains a new line character"
    return bool(run.text.endswith("\n") if run is not None else False)


def run_has_attr(run, attr):
    "check if a run has an attribute, for checking bold or italic for example"
    if not run:
        return None
    try:
        return getattr(run, attr)
    except AttributeError:
        # look at the font attribute if an AttributeError is thrown
        return getattr(run.font, attr)


def open_close_style(one_has_attr, two_has_attr, one_contains_break, output, attr):
    "open and close tags to include between two strings based on their attributes"
    open_tag, close_tag = html_open_close_tag(style=attr)
    if not open_tag or not close_tag:
        return output
    # add the close tag first
    if (
            (one_has_attr and one_contains_break) or
            (one_has_attr and two_has_attr is not True)):
        # check for new line
        if output.endswith("\n"):
            output = output.rstrip("\n") + close_tag + "\n"
        # check for font styles extending into whitespace
        elif output.endswith(" "):
            output = output.rstrip(" ") + close_tag + " "
        else:
            output += close_tag
    # add the open tag
    if (
            (two_has_attr and one_contains_break) or
            (two_has_attr and one_has_attr is not True)):
        output += open_tag
    return output


def run_open_close_style(run, prev_run, output, attr):
    "open and close tags to include between runs"
    # extract run object data for the more general purpose function
    return open_close_style(
        one_has_attr=run_has_attr(prev_run, attr),
        two_has_attr=run_has_attr(run, attr),
        one_contains_break=run_contains_break(prev_run),
        output=output,
        attr=attr
    )


def join_run_tags(run, prev_run, output=''):
    "process all the possible tags in the run"
    output = run_open_close_style(run, prev_run, output, 'italic')
    output = run_open_close_style(run, prev_run, output, 'bold')
    output = run_open_close_style(run, prev_run, output, 'subscript')
    output = run_open_close_style(run, prev_run, output, 'superscript')
    return output


def join_runs(runs):
    output = ''
    prev_run = None
    for run in runs:
        output = join_run_tags(run, prev_run, output)
        output += run.text
        prev_run = run
    # finish up by running one last time with prev_run
    output = join_run_tags('', prev_run, output)
    return output


if __name__ == "__main__":
    # debug while developing
    DIGEST_CONTENT = parse_content('tests/test_data/DIGEST 99999.docx')
    print(DIGEST_CONTENT.encode('utf-8'))

    """
    for file_name in ['DIGEST 20713.docx', 'DIGEST 24728.docx',
    'DIGEST 25783.docx', 'DIGEST 26726.docx']:
        print("\n\n")
        print(file_name)
        digest_content = parse_content(file_name)
        print(digest_content.encode('utf-8'))
    """
